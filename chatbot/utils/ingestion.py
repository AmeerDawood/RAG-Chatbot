"""
Document ingestion: text extraction, cleaning, chunking, and indexing into
the shared ChromaDB/LlamaIndex vector store.

Ported from the original project's chatbot/tasks.py with one correctness fix
(see ISSUES.md #3): every node now carries `document_id` in its metadata,
matching the key used by `delete_uploaded_file()` when it later asks Chroma
to remove that document's vectors. In the original project this metadata key
was never set, so deletions silently matched nothing and orphaned vectors
stayed searchable forever.
"""

import hashlib
import html
import logging
import os
import re

import docx
import fitz  # PyMuPDF
import pandas as pd
from django.conf import settings
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.readers import Document

from .embedding_setup import shared_index

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = settings.ALLOWED_UPLOAD_EXTENSIONS

_WHITESPACE_RE = re.compile(r'\s+')
_PAGE_MARKER_RE = re.compile(r'Page \d+ of \d+')

CHUNK_SIZE = 1024
CHUNK_OVERLAP = 200
SINGLE_NODE_CHAR_THRESHOLD = 1024


def clean_text(text, preserve_line_breaks=True):
    text = _PAGE_MARKER_RE.sub('', text)
    text = html.unescape(text)
    text = re.sub(r'[^\x00-\x7F؀-ۿ]+', ' ', text)

    if preserve_line_breaks:
        lines = [_WHITESPACE_RE.sub(' ', line) for line in text.splitlines()]
        text = '\n'.join(lines)
    else:
        text = _WHITESPACE_RE.sub(' ', text)

    return text.strip()


def generate_file_hash(filepath):
    hasher = hashlib.blake2b()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            hasher.update(chunk)
    return hasher.hexdigest()


def extract_text_from_file(filepath, ext):
    try:
        if ext == '.csv':
            df = pd.read_csv(filepath, encoding='utf-8', on_bad_lines='skip')
            df.drop_duplicates(inplace=True)
            df.columns = [col.strip().title().replace('_', ' ') for col in df.columns]
            df = df.map(lambda x: str(x).strip() if isinstance(x, str) else x)
            text = df.to_string(index=False)

        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        elif ext == '.pdf':
            with fitz.open(filepath) as pdf:
                text = ' '.join(page.get_text() for page in pdf)

        elif ext == '.docx':
            doc = docx.Document(filepath)
            text = ' '.join(para.text for para in doc.paragraphs)

        elif ext == '.xlsx':
            xls = pd.read_excel(filepath, sheet_name=None)
            all_dfs = [df.drop_duplicates().astype(str) for df in xls.values()]
            text = '\n'.join(df.to_string(index=False) for df in all_dfs)

        else:
            return None, f'Unsupported file extension: {ext}'

    except Exception as exc:
        return None, str(exc)

    return clean_text(text), None


def process_uploaded_file(uploaded_file_obj):
    """Extracts, chunks, embeds and indexes an UploadedFile. Updates its
    status field in place and returns a human-readable result message.
    """
    filepath = os.path.join(settings.MEDIA_ROOT, uploaded_file_obj.file.name)
    ext = os.path.splitext(filepath)[1].lower()

    logger.info('Processing file %s (ext=%s)', uploaded_file_obj.file.name, ext)

    if ext not in ALLOWED_EXTENSIONS:
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = 'Unsupported file format.'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return 'Unsupported file format.'

    file_hash = generate_file_hash(filepath)

    from .models import UploadedFile  # local import to avoid circularity at module load

    if UploadedFile.objects.filter(file_hash=file_hash, status='success').exclude(
        pk=uploaded_file_obj.pk
    ).exists():
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = 'This file has already been uploaded.'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return 'This file has already been uploaded.'

    cleaned_text, error = extract_text_from_file(filepath, ext)
    if error:
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = f'Failed to extract content: {error}'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return f'Failed to extract content: {error}'

    if not cleaned_text or not cleaned_text.strip():
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = 'The file appears to be empty.'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return 'The file appears to be empty.'

    word_count = len(cleaned_text.strip().split())
    unique_words = set(cleaned_text.lower().strip().split())
    if word_count < 5 or len(unique_words) < 3:
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = 'The file does not contain meaningful content.'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return 'The file does not contain meaningful content.'

    metadata = {
        'filename': uploaded_file_obj.file.name,
        'file_hash': file_hash,
        'file_type': ext,
        # Fix for ISSUES.md #3: this is the key delete_uploaded_file() filters on.
        'document_id': file_hash,
    }

    try:
        if len(cleaned_text) < SINGLE_NODE_CHAR_THRESHOLD:
            nodes = [Document(text=cleaned_text, metadata=metadata)]
        else:
            document = Document(text=cleaned_text, metadata=metadata)
            parser = SimpleNodeParser.from_defaults(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
            )
            nodes = parser.get_nodes_from_documents([document])

        shared_index.insert_nodes(nodes)
        logger.info('Indexed %d node(s) for %s', len(nodes), uploaded_file_obj.file.name)

        uploaded_file_obj.file_hash = file_hash
        uploaded_file_obj.status = 'success'
        uploaded_file_obj.error_message = None
        uploaded_file_obj.save(update_fields=['file_hash', 'status', 'error_message'])
        return 'File processed and indexed successfully.'

    except Exception as exc:
        logger.exception('Unexpected error during indexing of %s', uploaded_file_obj.file.name)
        uploaded_file_obj.status = 'failed'
        uploaded_file_obj.error_message = f'Unexpected error during indexing: {exc}'
        uploaded_file_obj.save(update_fields=['status', 'error_message'])
        return f'Unexpected error during indexing: {exc}'
